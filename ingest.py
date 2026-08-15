import os
import tempfile
from functools import lru_cache

import pandas as pd

from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings

from pinecone import Pinecone

from config import PINECONE_API_KEY, PINECONE_INDEX


# =========================================================
# CONFIGURATION
# =========================================================

MAX_FILES = 10

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".csv",
    ".xlsx",
    ".xls"
}


# =========================================================
# EMBEDDING MODEL
# =========================================================

@lru_cache(maxsize=1)
def get_embedding_model():

    print("Loading embedding model...")

    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )


# =========================================================
# PDF LOADER
# =========================================================

def load_pdf(file_path, filename):

    loader = PyPDFLoader(file_path)

    documents = loader.load()

    for document in documents:

        document.metadata["source"] = filename

    return documents


# =========================================================
# DOCX LOADER
# =========================================================

def load_docx(file_path, filename):

    doc = DocxDocument(file_path)

    text_parts = []

    # -------------------------
    # Paragraphs
    # -------------------------

    for paragraph in doc.paragraphs:

        text = paragraph.text.strip()

        if text:

            text_parts.append(text)


    # -------------------------
    # Tables
    # -------------------------

    for table in doc.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                cell_text = cell.text.strip()

                if cell_text:

                    row_text.append(cell_text)

            if row_text:

                text_parts.append(
                    " | ".join(row_text)
                )


    text = "\n".join(text_parts)


    if not text.strip():

        raise ValueError(
            "No readable text found in DOCX."
        )


    return [
        Document(
            page_content=text,
            metadata={
                "source": filename
            }
        )
    ]


# =========================================================
# TXT LOADER
# =========================================================

def load_txt(file_path, filename):

    with open(
        file_path,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as file:

        text = file.read()


    if not text.strip():

        raise ValueError(
            "TXT file is empty."
        )


    return [
        Document(
            page_content=text,
            metadata={
                "source": filename
            }
        )
    ]


# =========================================================
# CSV LOADER
# =========================================================

def load_csv(file_path, filename):

    dataframe = pd.read_csv(
        file_path
    )


    text = dataframe.to_string(
        index=False
    )


    if not text.strip():

        raise ValueError(
            "CSV file contains no readable data."
        )


    return [
        Document(
            page_content=text,
            metadata={
                "source": filename
            }
        )
    ]


# =========================================================
# EXCEL LOADER
# =========================================================

def load_excel(file_path, filename):

    excel_file = pd.ExcelFile(
        file_path
    )

    documents = []


    for sheet_name in excel_file.sheet_names:

        dataframe = pd.read_excel(
            file_path,
            sheet_name=sheet_name
        )


        text = dataframe.to_string(
            index=False
        )


        if text.strip():

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": filename,
                        "sheet": sheet_name
                    }
                )
            )


    if not documents:

        raise ValueError(
            "Excel file contains no readable data."
        )


    return documents


# =========================================================
# DOCUMENT LOADER SELECTOR
# =========================================================

def load_document(
    file_path,
    filename
):

    extension = os.path.splitext(
        filename
    )[1].lower()


    if extension == ".pdf":

        return load_pdf(
            file_path,
            filename
        )


    elif extension == ".docx":

        return load_docx(
            file_path,
            filename
        )


    elif extension == ".txt":

        return load_txt(
            file_path,
            filename
        )


    elif extension == ".csv":

        return load_csv(
            file_path,
            filename
        )


    elif extension in [
        ".xlsx",
        ".xls"
    ]:

        return load_excel(
            file_path,
            filename
        )


    else:

        raise ValueError(
            f"Unsupported file type: {extension}"
        )


# =========================================================
# INGEST DOCUMENTS
# =========================================================

def ingest_document(
    uploaded_files,
    namespace,
    progress_bar=None,
    status_text=None
):

    # -----------------------------------------------------
    # Validate file count
    # -----------------------------------------------------

    if not uploaded_files:

        raise ValueError(
            "No documents were uploaded."
        )


    if len(uploaded_files) > MAX_FILES:

        raise ValueError(
            f"Maximum {MAX_FILES} documents are allowed."
        )


    # -----------------------------------------------------
    # Connect to Pinecone
    # -----------------------------------------------------

    print(
        f"Connecting to Pinecone namespace: {namespace}"
    )


    pc = Pinecone(
        api_key=PINECONE_API_KEY
    )


    index = pc.Index(
        PINECONE_INDEX
    )


    # -----------------------------------------------------
    # Text splitter
    # -----------------------------------------------------

    splitter = RecursiveCharacterTextSplitter(

        chunk_size=500,

        chunk_overlap=100
    )


    # -----------------------------------------------------
    # Embedding model
    # -----------------------------------------------------

    embedding_model = get_embedding_model()


    # -----------------------------------------------------
    # Counters
    # -----------------------------------------------------

    processed_files = 0

    failed_files = []


    total_files = len(
        uploaded_files
    )


    # =====================================================
    # PROCESS EACH FILE
    # =====================================================

    for file_number, uploaded_file in enumerate(
        uploaded_files
    ):

        filename = uploaded_file.name

        temp_path = None


        # -------------------------------------------------
        # Progress
        # -------------------------------------------------

        if progress_bar:

            progress = (
                file_number / total_files
            )

            progress_bar.progress(
                int(progress * 100)
            )


        if status_text:

            status_text.text(
                f"Processing {filename} "
                f"({file_number + 1}/{total_files})..."
            )


        print(
            f"\nProcessing: {filename}"
        )


        try:

            # ---------------------------------------------
            # Check extension
            # ---------------------------------------------

            extension = os.path.splitext(
                filename
            )[1].lower()


            if extension not in SUPPORTED_EXTENSIONS:

                raise ValueError(
                    f"Unsupported file type: {extension}"
                )


            # ---------------------------------------------
            # Temporary file
            # ---------------------------------------------

            temp_file = tempfile.NamedTemporaryFile(
                delete=False,
                suffix=extension
            )


            temp_path = temp_file.name


            temp_file.write(
                uploaded_file.getbuffer()
            )


            temp_file.close()


            # ---------------------------------------------
            # Load document
            # ---------------------------------------------

            documents = load_document(
                temp_path,
                filename
            )


            if not documents:

                raise ValueError(
                    "No content could be extracted."
                )


            # ---------------------------------------------
            # Metadata
            # ---------------------------------------------

            for document in documents:

                document.metadata["source"] = (
                    filename
                )


            # ---------------------------------------------
            # Split documents
            # ---------------------------------------------

            print(
                f"Splitting {filename}..."
            )


            splits = splitter.split_documents(
                documents
            )


            if not splits:

                raise ValueError(
                    "No text chunks were created."
                )


            print(
                f"Created {len(splits)} chunks."
            )


            # ---------------------------------------------
            # Extract text
            # ---------------------------------------------

            texts = [

                document.page_content

                for document in splits

            ]


            # ---------------------------------------------
            # Create embeddings
            # ---------------------------------------------

            print(
                f"Creating embeddings for {filename}..."
            )


            embeddings = (
                embedding_model.embed_documents(
                    texts
                )
            )


            # ---------------------------------------------
            # Prepare vectors
            # ---------------------------------------------

            vectors = []


            for chunk_number, (
                document,
                embedding
            ) in enumerate(
                zip(
                    splits,
                    embeddings
                )
            ):

                metadata = {

                    "text":
                        document.page_content,

                    "source":
                        filename

                }


                # PDF page number

                if "page" in document.metadata:

                    metadata["page"] = (
                        document.metadata["page"]
                    )


                # Excel sheet

                if "sheet" in document.metadata:

                    metadata["sheet"] = (
                        document.metadata["sheet"]
                    )


                vector_id = (

                    f"{file_number}-"
                    f"{chunk_number}"

                )


                vectors.append({

                    "id":
                        vector_id,

                    "values":
                        embedding,

                    "metadata":
                        metadata

                })


            # ---------------------------------------------
            # Upload vectors
            # ---------------------------------------------

            print(
                f"Uploading {filename}..."
            )


            batch_size = 100


            for i in range(
                0,
                len(vectors),
                batch_size
            ):

                batch = vectors[
                    i:i + batch_size
                ]


                index.upsert(

                    vectors=batch,

                    namespace=namespace

                )


            print(
                f"SUCCESS: {filename}"
            )


            processed_files += 1


        except Exception as e:

            print(
                f"FAILED: {filename}: {e}"
            )


            failed_files.append(
                (
                    filename,
                    str(e)
                )
            )


        finally:

            # ---------------------------------------------
            # Delete temporary file
            # ---------------------------------------------

            if (
                temp_path
                and
                os.path.exists(temp_path)
            ):

                os.remove(
                    temp_path
                )


    # =====================================================
    # COMPLETE
    # =====================================================

    if progress_bar:

        progress_bar.progress(100)


    if status_text:

        status_text.text(
            "Document processing completed."
        )


    return {

        "processed_files":
            processed_files,

        "failed_files":
            failed_files

    }


# =========================================================
# DIRECT EXECUTION
# =========================================================

if __name__ == "__main__":

    print(
        "Run the application using:"
    )

    print(
        "streamlit run streamlit_app.py"
    )